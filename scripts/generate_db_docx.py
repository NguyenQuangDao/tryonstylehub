from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_section(document, title, purpose, fields):
    h = document.add_heading(title, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = document.add_paragraph(purpose)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = document.add_table(rows=1, cols=4)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Trường'
    hdr_cells[1].text = 'Kiểu dữ liệu'
    hdr_cells[2].text = 'Độ lớn'
    hdr_cells[3].text = 'Mô tả'
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for f in fields:
        row_cells = table.add_row().cells
        row_cells[0].text = f.get('Trường', '')
        row_cells[1].text = f.get('Kiểu dữ liệu', '')
        row_cells[2].text = f.get('Độ lớn', '')
        row_cells[3].text = f.get('Mô tả', '')

def build_data():
    return [
        {
            'title': 'Bảng User',
            'purpose': 'Mục đích: lưu thông tin người dùng, phân quyền và số dư token.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính, mặc định cuid()'},
                {'Trường': 'email', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Unique'},
                {'Trường': 'password', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'name', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'avatarUrl', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'role', 'Kiểu dữ liệu': 'enum', 'Độ lớn': '', 'Mô tả': 'ADMIN, USER, SELLER; Index; mặc định USER'},
                {'Trường': 'tokenBalance', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': 'Mặc định 0'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng Shop',
            'purpose': 'Mục đích: lưu cửa hàng thuộc người bán.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'name', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'slug', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Unique, Index'},
                {'Trường': 'logoUrl', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'bannerUrl', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'description', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'status', 'Kiểu dữ liệu': 'enum', 'Độ lớn': '', 'Mô tả': 'PENDING, ACTIVE, SUSPENDED; mặc định PENDING'},
                {'Trường': 'averageRating', 'Kiểu dữ liệu': 'float', 'Độ lớn': '', 'Mô tả': 'Mặc định 0'},
                {'Trường': 'totalSales', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': 'Mặc định 0'},
                {'Trường': 'ownerId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Unique; FK → User.id'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng Category',
            'purpose': 'Mục đích: quản lý danh mục sản phẩm (hỗ trợ phân cấp).',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'name', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'slug', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Unique, Index'},
                {'Trường': 'imageUrl', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'parentId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → Category.id; Index; có thể null'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng Product',
            'purpose': 'Mục đích: lưu thông tin sản phẩm.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'title', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'description', 'Kiểu dữ liệu': 'text', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'images', 'Kiểu dữ liệu': 'json', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'basePrice', 'Kiểu dữ liệu': 'decimal', 'Độ lớn': '10,2', 'Mô tả': ''},
                {'Trường': 'status', 'Kiểu dữ liệu': 'enum', 'Độ lớn': '', 'Mô tả': 'DRAFT, PUBLISHED, ARCHIVED; mặc định DRAFT'},
                {'Trường': 'shopId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → Shop.id; Index'},
                {'Trường': 'categoryId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → Category.id; Index'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng ProductVariant',
            'purpose': 'Mục đích: lưu biến thể của sản phẩm (kích thước, màu, v.v.).',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'productId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → Product.id; Index'},
                {'Trường': 'name', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'price', 'Kiểu dữ liệu': 'decimal', 'Độ lớn': '10,2', 'Mô tả': 'Có thể null'},
                {'Trường': 'stock', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': 'Mặc định 0'},
                {'Trường': 'sku', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Unique'},
                {'Trường': 'attributes', 'Kiểu dữ liệu': 'json', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng Order',
            'purpose': 'Mục đích: lưu đơn hàng.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'buyerId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → User.id; Index'},
                {'Trường': 'shopId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → Shop.id; Index'},
                {'Trường': 'totalAmount', 'Kiểu dữ liệu': 'decimal', 'Độ lớn': '10,2', 'Mô tả': ''},
                {'Trường': 'shippingAddress', 'Kiểu dữ liệu': 'json', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'status', 'Kiểu dữ liệu': 'enum', 'Độ lớn': '', 'Mô tả': 'PENDING, PROCESSING, SHIPPED, DELIVERED, CANCELLED'},
                {'Trường': 'paymentStatus', 'Kiểu dữ liệu': 'enum', 'Độ lớn': '', 'Mô tả': 'UNPAID, PAID, REFUNDED; mặc định UNPAID'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng OrderItem',
            'purpose': 'Mục đích: chi tiết sản phẩm trong đơn hàng.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'orderId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → Order.id; Index'},
                {'Trường': 'productId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → Product.id; Index'},
                {'Trường': 'productVariantId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → ProductVariant.id; Index; có thể null'},
                {'Trường': 'quantity', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'price', 'Kiểu dữ liệu': 'decimal', 'Độ lớn': '10,2', 'Mô tả': ''},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng Transaction',
            'purpose': 'Mục đích: lưu giao dịch ví/tài khoản của người dùng.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'userId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → User.id; Index'},
                {'Trường': 'amount', 'Kiểu dữ liệu': 'decimal', 'Độ lớn': '10,2', 'Mô tả': ''},
                {'Trường': 'type', 'Kiểu dữ liệu': 'enum', 'Độ lớn': '', 'Mô tả': 'DEPOSIT, PURCHASE, REFUND'},
                {'Trường': 'status', 'Kiểu dữ liệu': 'enum', 'Độ lớn': '', 'Mô tả': 'PENDING, SUCCESS, FAILED; Index; mặc định PENDING'},
                {'Trường': 'description', 'Kiểu dữ liệu': 'text', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng Review',
            'purpose': 'Mục đích: lưu đánh giá của người dùng cho sản phẩm.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'rating', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'comment', 'Kiểu dữ liệu': 'text', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'images', 'Kiểu dữ liệu': 'json', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'userId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → User.id; Index'},
                {'Trường': 'productId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → Product.id; Index'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng TokenPurchase',
            'purpose': 'Mục đích: lưu giao dịch mua token qua Stripe.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'userId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → User.id; Index'},
                {'Trường': 'stripePaymentId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Unique'},
                {'Trường': 'amount', 'Kiểu dữ liệu': 'decimal', 'Độ lớn': '10,2', 'Mô tả': ''},
                {'Trường': 'tokens', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'status', 'Kiểu dữ liệu': 'enum', 'Độ lớn': '', 'Mô tả': 'pending, completed, failed; mặc định pending'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
            ],
        },
        {
            'title': 'Bảng CostTracking',
            'purpose': 'Mục đích: theo dõi chi phí sử dụng dịch vụ.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'userId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → User.id; Index (ghép với service); có thể null'},
                {'Trường': 'service', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'operation', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'cost', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': 'Mặc định 0'},
                {'Trường': 'details', 'Kiểu dữ liệu': 'text', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now(); Index tổng hợp (userId, service)'},
            ],
        },
        {
            'title': 'Bảng VirtualModel',
            'purpose': 'Mục đích: lưu cấu hình avatar/mô hình ảo của người dùng.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'userId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → User.id; Index; có thể null'},
                {'Trường': 'sessionId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Index; có thể null'},
                {'Trường': 'avatarName', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'avatarImage', 'Kiểu dữ liệu': 'longtext', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'isPublic', 'Kiểu dữ liệu': 'boolean', 'Độ lớn': '', 'Mô tả': 'Mặc định false'},
                {'Trường': 'height', 'Kiểu dữ liệu': 'float', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'weight', 'Kiểu dữ liệu': 'float', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'gender', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'hairColor', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'hairStyle', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'bodyShape', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'skinTone', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'eyeColor', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'faceShape', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'beardStyle', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'tattoos', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'piercings', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'clothingStyle', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'accessories', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'footwearType', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'colorPalette', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'ageAppearance', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'bodyProportionPreset', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Có thể null'},
                {'Trường': 'muscleLevel', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'fatLevel', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'shoulderWidth', 'Kiểu dữ liệu': 'float', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'waistSize', 'Kiểu dữ liệu': 'float', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'hipSize', 'Kiểu dữ liệu': 'float', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'legLength', 'Kiểu dữ liệu': 'float', 'Độ lớn': '', 'Mô tả': 'Có thể null'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng BlogPost',
            'purpose': 'Mục đích: bài viết blog của người dùng.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'authorId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → User.id; Index'},
                {'Trường': 'title', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': ''},
                {'Trường': 'content', 'Kiểu dữ liệu': 'longtext', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'media', 'Kiểu dữ liệu': 'json', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'category', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Index; có thể null'},
                {'Trường': 'tags', 'Kiểu dữ liệu': 'json', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'status', 'Kiểu dữ liệu': 'enum', 'Độ lớn': '', 'Mô tả': 'DRAFT, PUBLISHED, ARCHIVED; Index; mặc định PUBLISHED'},
                {'Trường': 'likesCount', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': 'Mặc định 0'},
                {'Trường': 'savesCount', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': 'Mặc định 0'},
                {'Trường': 'commentsCount', 'Kiểu dữ liệu': 'int', 'Độ lớn': '', 'Mô tả': 'Mặc định 0'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng BlogComment',
            'purpose': 'Mục đích: bình luận cho bài viết blog.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'postId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → BlogPost.id; Index'},
                {'Trường': 'userId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → User.id; Index'},
                {'Trường': 'content', 'Kiểu dữ liệu': 'longtext', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'media', 'Kiểu dữ liệu': 'json', 'Độ lớn': '', 'Mô tả': ''},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': 'updatedAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Tự động cập nhật'},
            ],
        },
        {
            'title': 'Bảng BlogLike',
            'purpose': 'Mục đích: lưu lượt thích cho bài viết.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'postId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → BlogPost.id'},
                {'Trường': 'userId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → User.id; Index'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': '(ràng buộc)', 'Kiểu dữ liệu': '', 'Độ lớn': '', 'Mô tả': 'UNIQUE(postId, userId)'},
            ],
        },
        {
            'title': 'Bảng BlogSave',
            'purpose': 'Mục đích: lưu lượt lưu bài viết.',
            'fields': [
                {'Trường': 'id', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'Khóa chính'},
                {'Trường': 'postId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → BlogPost.id'},
                {'Trường': 'userId', 'Kiểu dữ liệu': 'varchar', 'Độ lớn': '255', 'Mô tả': 'FK → User.id; Index'},
                {'Trường': 'createdAt', 'Kiểu dữ liệu': 'datetime', 'Độ lớn': '', 'Mô tả': 'Mặc định now()'},
                {'Trường': '(ràng buộc)', 'Kiểu dữ liệu': '', 'Độ lớn': '', 'Mô tả': 'UNIQUE(postId, userId)'},
            ],
        },
    ]

def main():
    document = Document()
    title = document.add_heading('Tài liệu các bảng (MySQL) dựa trên schema Prisma', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note = document.add_paragraph('Nguồn: prisma/schema.prisma và docs/pttkht/diagrams/structure/DB_Diagram.puml.')
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for section in build_data():
        add_section(document, section['title'], section['purpose'], section['fields'])
    for p in document.paragraphs:
        for run in p.runs:
            run.font.size = Pt(11)
    document.save('docs/pttkht/diagrams/structure/DB_Table_Documentation.docx')

if __name__ == '__main__':
    main()
